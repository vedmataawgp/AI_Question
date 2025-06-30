import os
import csv
import logging
from typing import List, Dict, Tuple
from werkzeug.utils import secure_filename
import pandas as pd

# Import document processing libraries with fallbacks
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import win32com.client
    DOC_AVAILABLE = True
except ImportError:
    DOC_AVAILABLE = False

class FileProcessor:
    """File processor for handling bulk content import"""
    
    ALLOWED_EXTENSIONS = {'txt', 'csv', 'docx', 'doc'}
    
    def __init__(self):
        self.logger = logging.getLogger(__name__)
    
    def is_allowed_file(self, filename: str) -> bool:
        """Check if file extension is allowed"""
        return '.' in filename and \
               filename.rsplit('.', 1)[1].lower() in self.ALLOWED_EXTENSIONS
    
    def get_file_type(self, filename: str) -> str:
        """Get file type from filename"""
        if '.' in filename:
            return filename.rsplit('.', 1)[1].lower()
        return 'unknown'
    
    def process_file(self, file_path: str, file_type: str) -> Tuple[List[Dict], str]:
        """Process uploaded file and extract content items"""
        try:
            if file_type == 'txt':
                return self._process_txt_file(file_path)
            elif file_type == 'csv':
                return self._process_csv_file(file_path)
            elif file_type == 'docx':
                return self._process_docx_file(file_path)
            elif file_type == 'doc':
                return self._process_doc_file(file_path)
            else:
                return [], f"Unsupported file type: {file_type}"
        
        except Exception as e:
            self.logger.error(f"Error processing file {file_path}: {e}")
            return [], f"Error processing file: {str(e)}"
    
    def _process_txt_file(self, file_path: str) -> Tuple[List[Dict], str]:
        """Process TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
            
            if not content:
                return [], "File is empty"
            
            # Split content into paragraphs or sections
            paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
            
            items = []
            for i, paragraph in enumerate(paragraphs, 1):
                if len(paragraph) > 10:  # Skip very short paragraphs
                    # Try to extract title from first line
                    lines = paragraph.split('\n')
                    title = lines[0] if len(lines[0]) < 100 else f"Content Item {i}"
                    content_text = paragraph if len(lines) == 1 else '\n'.join(lines[1:]) or paragraph
                    
                    items.append({
                        'title': title,
                        'content': content_text
                    })
            
            return items, ""
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    content = file.read().strip()
                
                paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
                items = []
                for i, paragraph in enumerate(paragraphs, 1):
                    if len(paragraph) > 10:
                        lines = paragraph.split('\n')
                        title = lines[0] if len(lines[0]) < 100 else f"Content Item {i}"
                        content_text = paragraph if len(lines) == 1 else '\n'.join(lines[1:]) or paragraph
                        
                        items.append({
                            'title': title,
                            'content': content_text
                        })
                
                return items, ""
                
            except Exception as e:
                return [], f"Error reading text file: {str(e)}"
    
    def _process_csv_file(self, file_path: str) -> Tuple[List[Dict], str]:
        """Process CSV file"""
        try:
            # Try to read with pandas first
            df = pd.read_csv(file_path, encoding='utf-8')
            
            # Expected columns: title, content (or question, answer)
            items = []
            
            if 'title' in df.columns and 'content' in df.columns:
                for _, row in df.iterrows():
                    if pd.notna(row['title']) and pd.notna(row['content']):
                        items.append({
                            'title': str(row['title']).strip(),
                            'content': str(row['content']).strip()
                        })
            
            elif 'question' in df.columns and 'answer' in df.columns:
                for _, row in df.iterrows():
                    if pd.notna(row['question']) and pd.notna(row['answer']):
                        items.append({
                            'title': str(row['question']).strip(),
                            'content': str(row['answer']).strip()
                        })
            
            else:
                # Try to use first two columns
                if len(df.columns) >= 2:
                    col1, col2 = df.columns[0], df.columns[1]
                    for _, row in df.iterrows():
                        if pd.notna(row[col1]) and pd.notna(row[col2]):
                            items.append({
                                'title': str(row[col1]).strip(),
                                'content': str(row[col2]).strip()
                            })
                else:
                    return [], "CSV file must have at least 2 columns (title/content or question/answer)"
            
            return items, ""
            
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                df = pd.read_csv(file_path, encoding='latin-1')
                items = []
                
                if 'title' in df.columns and 'content' in df.columns:
                    for _, row in df.iterrows():
                        if pd.notna(row['title']) and pd.notna(row['content']):
                            items.append({
                                'title': str(row['title']).strip(),
                                'content': str(row['content']).strip()
                            })
                elif len(df.columns) >= 2:
                    col1, col2 = df.columns[0], df.columns[1]
                    for _, row in df.iterrows():
                        if pd.notna(row[col1]) and pd.notna(row[col2]):
                            items.append({
                                'title': str(row[col1]).strip(),
                                'content': str(row[col2]).strip()
                            })
                
                return items, ""
                
            except Exception as e:
                return [], f"Error reading CSV file: {str(e)}"
        
        except Exception as e:
            return [], f"Error processing CSV file: {str(e)}"
    
    def _process_docx_file(self, file_path: str) -> Tuple[List[Dict], str]:
        """Process DOCX file"""
        if not DOCX_AVAILABLE:
            return [], "python-docx library not available. Install it to process DOCX files."
        
        try:
            doc = DocxDocument(file_path)
            
            items = []
            current_title = ""
            current_content = []
            
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if not text:
                    continue
                
                # Check if this might be a heading (simple heuristic)
                if (len(text) < 100 and 
                    (paragraph.style.name.startswith('Heading') or 
                     text.isupper() or 
                     (len(current_content) > 0 and text.endswith(':')))):
                    
                    # Save previous section if exists
                    if current_title and current_content:
                        items.append({
                            'title': current_title,
                            'content': '\n'.join(current_content)
                        })
                    
                    # Start new section
                    current_title = text
                    current_content = []
                else:
                    current_content.append(text)
            
            # Add last section
            if current_title and current_content:
                items.append({
                    'title': current_title,
                    'content': '\n'.join(current_content)
                })
            
            # If no clear structure found, treat as single document
            if not items:
                all_text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
                if all_text.strip():
                    items.append({
                        'title': 'Document Content',
                        'content': all_text.strip()
                    })
            
            return items, ""
            
        except Exception as e:
            return [], f"Error processing DOCX file: {str(e)}"
    
    def _process_doc_file(self, file_path: str) -> Tuple[List[Dict], str]:
        """Process DOC file"""
        # DOC file processing is complex and requires Word application on Windows
        # For now, return an error message suggesting conversion to DOCX
        return [], "DOC file processing requires Microsoft Word. Please convert to DOCX format and try again."
    
    def save_uploaded_file(self, file, upload_folder: str) -> Tuple[str, str]:
        """Save uploaded file and return filename and full path"""
        try:
            filename = secure_filename(file.filename)
            if not filename:
                return "", ""
            
            # Generate unique filename to avoid conflicts
            import time
            timestamp = str(int(time.time()))
            name, ext = os.path.splitext(filename)
            unique_filename = f"{name}_{timestamp}{ext}"
            
            file_path = os.path.join(upload_folder, unique_filename)
            file.save(file_path)
            
            return unique_filename, file_path
            
        except Exception as e:
            logging.error(f"Error saving uploaded file: {e}")
            return "", ""
